import os
import hashlib
import mimetypes
import magic
from PIL import Image
from pathlib import Path
import logging
from datetime import datetime, timedelta


def calculate_file_hash(file_path):
    """Calculate SHA-256 hash of a file"""
    hash_sha256 = hashlib.sha256()
    try:
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    except Exception as e:
        logging.error(f"Error calculating hash for {file_path}: {e}")
        return None


def get_file_mime_type(file_path):
    """Get MIME type of a file using python-magic and mimetypes"""
    try:
        # Try python-magic first (more accurate)
        mime_type = magic.from_file(file_path, mime=True)
        if mime_type:
            return mime_type
    except:
        pass
    
    # Fallback to mimetypes
    mime_type, _ = mimetypes.guess_type(file_path)
    return mime_type or 'application/octet-stream'


def get_image_dimensions(file_path):
    """Get dimensions of an image file"""
    try:
        with Image.open(file_path) as img:
            return f"{img.width}x{img.height}"
    except Exception as e:
        logging.error(f"Error getting image dimensions for {file_path}: {e}")
        return None


def count_document_pages(file_path):
    """Count pages in a document (PDF, Word, etc.)"""
    try:
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                return len(reader.pages)
        
        elif file_extension in ['.docx', '.doc']:
            from docx import Document
            doc = Document(file_path)
            # Approximate page count based on paragraphs
            return max(1, len(doc.paragraphs) // 10)
        
        elif file_extension in ['.pptx', '.ppt']:
            from pptx import Presentation
            prs = Presentation(file_path)
            return len(prs.slides)
        
    except Exception as e:
        logging.error(f"Error counting pages for {file_path}: {e}")
    
    return None


def count_words_in_text(file_path):
    """Count words in a text file"""
    try:
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                return len(content.split())
        
        elif file_extension in ['.docx', '.doc']:
            from docx import Document
            doc = Document(file_path)
            word_count = 0
            for paragraph in doc.paragraphs:
                word_count += len(paragraph.text.split())
            return word_count
        
    except Exception as e:
        logging.error(f"Error counting words for {file_path}: {e}")
    
    return None


def cleanup_old_files(folder_path, max_age_hours=24):
    """Clean up files older than specified hours"""
    try:
        current_time = datetime.now()
        max_age = timedelta(hours=max_age_hours)
        
        for file_path in Path(folder_path).iterdir():
            if file_path.is_file():
                file_age = current_time - datetime.fromtimestamp(file_path.stat().st_mtime)
                if file_age > max_age:
                    file_path.unlink()
                    logging.info(f"Cleaned up old file: {file_path}")
        
    except Exception as e:
        logging.error(f"Error during cleanup of {folder_path}: {e}")


def format_file_size(size_bytes):
    """Format file size in human readable format"""
    if size_bytes == 0:
        return "0 B"
    
    size_names = ["B", "KB", "MB", "GB", "TB"]
    i = 0
    while size_bytes >= 1024 and i < len(size_names) - 1:
        size_bytes /= 1024.0
        i += 1
    
    return f"{size_bytes:.1f} {size_names[i]}"


def validate_file_security(file_path):
    """Perform security validation on uploaded files"""
    try:
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > 50 * 1024 * 1024:  # 50MB limit
            return False, "File size exceeds 50MB limit"
        
        # Check MIME type against file extension
        mime_type = get_file_mime_type(file_path)
        file_extension = Path(file_path).suffix.lower()
        
        # Basic MIME type validation
        valid_mimes = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel',
            '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
            '.ppt': 'application/vnd.ms-powerpoint',
            '.txt': 'text/plain',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
        }
        
        expected_mime = valid_mimes.get(file_extension)
        if expected_mime and not mime_type.startswith(expected_mime.split('/')[0]):
            return False, f"File content doesn't match extension {file_extension}"
        
        return True, "File validation passed"
        
    except Exception as e:
        logging.error(f"Error validating file {file_path}: {e}")
        return False, "File validation failed"


def generate_secure_filename(original_filename, file_id):
    """Generate a secure filename with UUID and sanitized original name"""
    # Sanitize original filename
    safe_chars = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789.-_"
    sanitized = ''.join(c for c in original_filename if c in safe_chars)
    
    # Get file extension
    extension = Path(original_filename).suffix.lower()
    
    # Combine with UUID
    return f"{file_id}_{sanitized[:50]}{extension}"


class ConversionProgress:
    """Track conversion progress for real-time updates"""
    
    def __init__(self, file_id):
        self.file_id = file_id
        self.progress = 0
        self.status = "initializing"
        self.message = "Starting conversion..."
        self.start_time = datetime.now()
    
    def update(self, progress, status, message):
        self.progress = progress
        self.status = status
        self.message = message
    
    def complete(self, success=True, message="Conversion completed"):
        self.progress = 100
        self.status = "completed" if success else "failed"
        self.message = message
        self.end_time = datetime.now()
    
    def get_duration(self):
        if hasattr(self, 'end_time'):
            return (self.end_time - self.start_time).total_seconds()
        return (datetime.now() - self.start_time).total_seconds()
    
    def to_dict(self):
        return {
            'file_id': self.file_id,
            'progress': self.progress,
            'status': self.status,
            'message': self.message,
            'duration': self.get_duration()
        }


# Global progress tracker
conversion_progress = {}


def get_conversion_progress(file_id):
    """Get conversion progress for a file"""
    return conversion_progress.get(file_id)


def set_conversion_progress(file_id, progress, status, message):
    """Set conversion progress for a file"""
    if file_id not in conversion_progress:
        conversion_progress[file_id] = ConversionProgress(file_id)
    
    conversion_progress[file_id].update(progress, status, message)